import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock
from langchain_core.documents import Document
from book_see_rag.ingestion.loader import load_document
from book_see_rag.ingestion.splitter import is_noisy_chunk, normalize_chunk_text, split_documents


# ── loader ────────────────────────────────────────────────────

def test_load_txt(tmp_path):
    f = tmp_path / "test.txt"
    f.write_text("Hello\nWorld", encoding="utf-8")
    docs = load_document(str(f))
    assert len(docs) == 1
    assert "Hello" in docs[0].page_content


def test_load_md(tmp_path):
    f = tmp_path / "test.md"
    f.write_text("# Title\n\nContent here", encoding="utf-8")
    docs = load_document(str(f))
    assert len(docs) == 1
    assert "Title" in docs[0].page_content


def test_load_unsupported_format(tmp_path):
    f = tmp_path / "test.xlsx"
    f.write_bytes(b"fake data")
    with pytest.raises(ValueError, match="不支持的文件格式"):
        load_document(str(f))


def test_load_docx(tmp_path):
    from docx import Document as DocxDocument
    f = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("第一段内容")
    doc.add_paragraph("第二段内容")
    doc.save(str(f))

    docs = load_document(str(f))
    assert len(docs) == 1
    assert "第一段内容" in docs[0].page_content


# ── splitter ──────────────────────────────────────────────────

def test_split_documents_produces_chunks():
    long_text = "这是一段很长的中文文本。" * 100
    docs = [Document(page_content=long_text, metadata={"source": "test.txt"})]
    chunks = split_documents(docs)
    assert len(chunks) > 1
    for c in chunks:
        assert len(c.page_content) <= 600  # chunk_size=512 + 一定余量


def test_split_preserves_metadata():
    docs = [Document(page_content="内容" * 50, metadata={"source": "a.txt", "page": 1})]
    chunks = split_documents(docs)
    for c in chunks:
        assert c.metadata["source"] == "a.txt"


def test_normalize_chunk_text_removes_blank_noise_lines():
    text = "第一行 \n\n  \n 第二行   内容\n"
    assert normalize_chunk_text(text) == "第一行\n第二行 内容"


def test_is_noisy_chunk_detects_repeated_noise():
    assert is_noisy_chunk("V V V V V V V V V")
    assert is_noisy_chunk("熟悉 F F F F F F F")
    assert not is_noisy_chunk("具有 7 年 Python 后端开发经验，熟悉 RAG 和模型部署。")


def test_split_documents_drops_noisy_chunks():
    docs = [Document(page_content=("正常内容。 " * 40) + ("\nV V V V V V V V V\n" * 20), metadata={"source": "a.txt"})]
    chunks = split_documents(docs)
    assert chunks
    assert all("V V V V V" not in c.page_content for c in chunks)


# ── pdf loader ────────────────────────────────────────────────

def test_pdf_loader_text_version(tmp_path):
    with patch("book_see_rag.ingestion.pdf_loader.pdfplumber") as mock_pdf:
        page_mock = MagicMock()
        page_mock.extract_text.return_value = "这是PDF文字内容 " * 20
        mock_pdf.open.return_value.__enter__.return_value.pages = [page_mock]

        from book_see_rag.ingestion.pdf_loader import load_pdf
        docs = load_pdf("fake.pdf")
        assert len(docs) == 1
        assert docs[0].metadata["is_ocr"] is False


def test_pdf_loader_triggers_ocr_for_sparse_pages():
    with patch("book_see_rag.ingestion.pdf_loader.pdfplumber") as mock_pdf, \
         patch("book_see_rag.ingestion.pdf_loader._ocr_with_marker") as mock_ocr:

        page_mock = MagicMock()
        page_mock.extract_text.return_value = "少"  # < 50 chars，触发 OCR
        mock_pdf.open.return_value.__enter__.return_value.pages = [page_mock] * 3
        mock_ocr.return_value = [Document(page_content="OCR结果", metadata={"is_ocr": True})]

        from book_see_rag.ingestion.pdf_loader import load_pdf
        docs = load_pdf("scan.pdf")
        mock_ocr.assert_called_once()
        assert docs[0].metadata["is_ocr"] is True
