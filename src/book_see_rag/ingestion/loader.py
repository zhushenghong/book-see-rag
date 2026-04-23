from pathlib import Path
import logging
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def load_document(path: str) -> list[Document]:
    """
    根据文件扩展名路由到对应加载器。
    支持：.pdf / .docx / .txt / .md
    """
    suffix = Path(path).suffix.lower()
    logger.info("Loading document path=%s suffix=%s", path, suffix)

    match suffix:
        case ".pdf":
            from book_see_rag.ingestion.pdf_loader import load_pdf
            return load_pdf(path)

        case ".docx":
            return _load_docx(path)

        case ".txt" | ".md" | ".markdown":
            return _load_text(path)

        case _:
            raise ValueError(f"不支持的文件格式：{suffix}，可选：pdf / docx / txt / md")


def _load_docx(path: str) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    logger.info("Loaded DOCX path=%s paragraphs=%s chars=%s", path, len(paragraphs), len(full_text))
    return [Document(page_content=full_text, metadata={"source": path, "is_ocr": False})]


def _load_text(path: str) -> list[Document]:
    text = Path(path).read_text(encoding="utf-8")
    logger.info("Loaded text path=%s chars=%s", path, len(text))
    return [Document(page_content=text, metadata={"source": path, "is_ocr": False})]
